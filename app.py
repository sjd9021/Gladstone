"""
Gladstone survey photo-sheet builder.

Upload all survey photos at once, type a caption under each, and download a
photo sheet formatted for the destination report:

  * WK Webster  -> photo block and caption table RIGHT aligned
  * Gladstone   -> photo block and caption table CENTRE aligned

Both use Arial 11 written as *direct* run formatting (not just the Normal
style), so that pasting into a report whose Normal style is Times New Roman
does not silently re-font the captions.

Layout constants were measured from real reports in the mail archive
(G-2175-25GW, G-1754-25BW, G-1222-25BW for WKW; G-1623-25G, G-1544-25B for
Gladstone).
"""

import io
import logging

import docx
import streamlit as st
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageOps

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------- constants

FONT_NAME = "Arial"
FONT_SIZE_PT = 11

IMG_W_CM = 7.14          # measured: every photo in every sample report
IMG_H_CM = 5.24
CAPTION_ROW_H_CM = 1.0

# The photos in a row are separated by a three-space run (as the existing
# sheets do). Arial's space advance is 0.2778 em, so at 11pt three spaces are
# 3 * 0.2778 * 11pt = 0.3233 cm. The caption table is widened by exactly that
# much so its two cells stay centred under the two photos; the real reports
# fudge this with a wider second column (gridCol 4048/4320) and end up ~1.6 mm
# out. Unlike centre alignment, a right-aligned row shows the error in full.
SPACER_TEXT = "   "
SPACER_CM = len(SPACER_TEXT) * 0.2778 * FONT_SIZE_PT / 72 * 2.54

# Word's default table cell margin, 108 twips. Applied as a paragraph indent
# inside caption cells rather than as a cell margin - see _fix_table_layout.
CELL_MARGIN_CM = 108 / 1440 * 2.54

# Page setup per destination, so the sheet previews how the block will sit
# in the real report. Only affects the sheet itself; a copy/paste carries
# the content, not the section properties.
LAYOUTS = {
    "wkw": {
        "label": "WK Webster",
        "para_align": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "table_align": WD_TABLE_ALIGNMENT.RIGHT,
        "html_align": "right",
        "align_word": "right",
        "page_w_cm": 21.0, "page_h_cm": 29.7,       # A4
        "margin_lr_cm": 1.9, "margin_t_cm": 2.54, "margin_b_cm": 1.27,
    },
    "gladstone": {
        "label": "Gladstone",
        "para_align": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "table_align": WD_TABLE_ALIGNMENT.CENTER,
        "html_align": "center",
        "align_word": "centre",
        "page_w_cm": 21.0, "page_h_cm": 29.7,       # A4
        "margin_lr_cm": 2.54, "margin_t_cm": 2.54, "margin_b_cm": 1.27,
    },
}

VALID_IMAGE_TYPES = {
    "image/jpeg", "image/jpg", "image/png", "image/gif",
    "image/bmp", "image/tiff", "image/webp",
}
VALID_PDF_TYPES = {"application/pdf"}

UPLOAD_EXTS = ["jpg", "jpeg", "png", "gif", "bmp", "tiff", "webp", "pdf"]


# ------------------------------------------------------------ image loading

def process_uploaded_image(uploaded_file):
    """Normalise an upload to JPEG bytes that python-docx will accept.

    Handles CMYK/RGBA/palette modes, honours EXIF orientation (phone photos
    are routinely stored rotated), and rasterises the first page of a PDF.
    Returns raw JPEG bytes, or None on failure.
    """
    if uploaded_file is None:
        return None

    ftype = (uploaded_file.type or "").lower()
    name = uploaded_file.name

    if ftype not in VALID_IMAGE_TYPES | VALID_PDF_TYPES:
        st.error(f"{name}: unsupported file type ({ftype or 'unknown'}).")
        return None

    try:
        uploaded_file.seek(0)

        if ftype in VALID_PDF_TYPES:
            try:
                from pdf2image import convert_from_bytes
            except ImportError:
                st.error(f"{name}: PDF support unavailable on this server.")
                return None
            pages = convert_from_bytes(
                uploaded_file.read(), dpi=200, first_page=1, last_page=1
            )
            if not pages:
                st.error(f"{name}: could not read any page from the PDF.")
                return None
            pil_image = pages[0]
        else:
            pil_image = Image.open(uploaded_file)
            pil_image = ImageOps.exif_transpose(pil_image)

        if pil_image.mode not in ("RGB", "L"):
            pil_image = pil_image.convert("RGB")

        buf = io.BytesIO()
        pil_image.save(buf, format="JPEG", quality=90)
        return buf.getvalue()

    except Exception as exc:                       # noqa: BLE001 - surfaced to user
        logger.exception("failed to process %s", name)
        st.error(f"{name}: could not be processed ({exc}).")
        return None


# --------------------------------------------------------- docx formatting

def _force_font(run):
    """Write Arial 11 directly onto the run.

    python-docx's ``run.font.name`` only sets w:ascii and w:hAnsi. Word uses
    w:cs for complex-script and w:eastAsia for CJK runs, so a caption can
    still drift on paste unless all four are set. Direct run formatting
    survives 'use destination styles', which style-level Arial does not.
    """
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PT)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)


def _clear_borders(table):
    """Blank every cell border, matching the existing hand-made sheets."""
    for cell in table._tbl.iter_tcs():
        tcPr = cell.tcPr
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "nil")
            borders.append(el)
        tcPr.append(borders)


def _caption_widths(n_photos):
    """Column widths that put one cell exactly under each photo.

    The photo row is  [photo][spacer][photo], so the caption table gets the
    same alternating columns. The narrow spacer column is empty and borderless;
    it exists so each caption column sits exactly under its photo rather than
    the pair being split evenly. Total table width equals total photo-row
    width, so both edges of the block line up.
    """
    widths = [IMG_W_CM]
    for _ in range(n_photos - 1):
        widths += [SPACER_CM, IMG_W_CM]
    return widths


def _fix_table_layout(table, jc_val, widths_cm):
    """Pin the caption table the way the real reports do it.

    Without ``tblLayout fixed`` Word and LibreOffice widen the cell to fit the
    caption on one line, so a long caption spills past the photo instead of
    wrapping under it. The real sheets also repeat ``jc`` on the row, which is
    what keeps the block flush right once it lands in the report.
    """
    twips = [int(round(w / 2.54 * 1440)) for w in widths_cm]

    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for col, w in zip(grid.findall(qn("w:gridCol")), twips):
            col.set(qn("w:w"), str(w))

    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        jc = trPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            trPr.append(jc)
        jc.set(qn("w:val"), jc_val)
        for idx, (cell, w) in enumerate(zip(row.cells, twips)):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")

            # Every column gets zero cell margins, for two separate reasons.
            # The narrow spacer columns are thinner than Word's default 108
            # twip margins, and Word refuses to draw a cell narrower than its
            # own margins - it widens the column and pushes the captions to its
            # right out of line. And for a right-aligned table Word offsets the
            # whole table by the trailing cell margin so that cell *content*,
            # not the cell edge, meets the page margin, which shifts every
            # caption 0.19 cm right of its photo. With the margins at zero the
            # table's edges are the photo block's edges. The same 0.19 cm inset
            # is restored as a paragraph indent inside the caption cells, so
            # captions still wrap at the width the existing reports use.
            tcMar = OxmlElement("w:tcMar")
            for side in ("left", "right"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:w"), "0")
                el.set(qn("w:type"), "dxa")
                tcMar.append(el)
            tcPr.append(tcMar)


def _add_photo_row(document, chunk, layout):
    """One row of photos (1 or 2) followed by its caption table."""
    para = document.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.alignment = layout["para_align"]
    run = para.add_run()

    for idx, (_, jpeg) in enumerate(chunk):
        if idx:
            run.add_text(SPACER_TEXT)       # between photos only
        run.add_picture(io.BytesIO(jpeg), width=Cm(IMG_W_CM), height=Cm(IMG_H_CM))
    _force_font(run)

    widths = _caption_widths(len(chunk))

    table = document.add_table(rows=1, cols=len(widths))
    table.autofit = False
    table.alignment = layout["table_align"]

    for idx, (caption, _) in enumerate(chunk):
        cell = table.cell(0, idx * 2)          # skip the spacer columns
        cell_para = cell.paragraphs[0]
        cell_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   # centred in both formats
        cell_para.paragraph_format.left_indent = Cm(CELL_MARGIN_CM)
        cell_para.paragraph_format.right_indent = Cm(CELL_MARGIN_CM)
        _force_font(cell_para.add_run(caption or ""))

    for row in table.rows:
        row.height = Cm(CAPTION_ROW_H_CM)

    _clear_borders(table)
    _fix_table_layout(table, layout["html_align"], widths)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def build_docx(items, layout_key):
    """items: list of (caption, jpeg_bytes). Returns a .docx as bytes."""
    layout = LAYOUTS[layout_key]

    document = docx.Document()          # per-call, never module-level
    normal = document.styles["Normal"].font
    normal.name = FONT_NAME
    normal.size = Pt(FONT_SIZE_PT)

    section = document.sections[0]
    section.page_width = Cm(layout["page_w_cm"])
    section.page_height = Cm(layout["page_h_cm"])
    section.left_margin = Cm(layout["margin_lr_cm"])
    section.right_margin = Cm(layout["margin_lr_cm"])
    section.top_margin = Cm(layout["margin_t_cm"])
    section.bottom_margin = Cm(layout["margin_b_cm"])

    for i in range(0, len(items), 2):
        _add_photo_row(document, items[i:i + 2], layout)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ----------------------------------------------------------------- the app

def _file_key(f):
    return f"{f.name}::{f.size}"


def _load_items(files, order):
    """Return [(caption, jpeg_bytes)] in display order, caching decoded JPEGs."""
    cache = st.session_state.setdefault("jpeg_cache", {})
    by_key = {_file_key(f): f for f in files}
    items = []
    for key in order:
        f = by_key.get(key)
        if f is None:
            continue
        if key not in cache:
            cache[key] = process_uploaded_image(f)
        if cache[key] is None:
            continue
        items.append((st.session_state.get(f"cap::{key}", "") or "", cache[key]))
    return items


def main():
    st.set_page_config(page_title="Gladstone Photo Sheet", layout="centered")
    st.markdown(
        "<h1 style='text-align:center;color:grey;'>Survey Photo Sheet</h1>",
        unsafe_allow_html=True,
    )

    files = st.file_uploader(
        "Drop all survey photos here (you can select many at once)",
        type=UPLOAD_EXTS,
        accept_multiple_files=True,
        help="JPEG, PNG, GIF, BMP, TIFF, WEBP or PDF (first page of a PDF is used).",
    )

    if not files:
        st.info("Upload photos to begin. Two photos go side by side per row.")
        return

    keys = [_file_key(f) for f in files]
    order = [k for k in st.session_state.get("order", []) if k in keys]
    order += [k for k in keys if k not in order]
    st.session_state["order"] = order

    by_key = {_file_key(f): f for f in files}

    st.caption(f"{len(order)} photo(s). Type the description under each; use the arrows to reorder.")

    for pos, key in enumerate(order):
        f = by_key[key]
        col_img, col_txt, col_up, col_dn = st.columns([2, 6, 1, 1])
        with col_img:
            if (f.type or "").lower() in VALID_PDF_TYPES:
                st.write("PDF")
            else:
                f.seek(0)
                st.image(f, width=110)
        with col_txt:
            st.text_input(
                f"Photo {pos + 1} description",
                key=f"cap::{key}",
                label_visibility="collapsed",
                placeholder=f"Photo {pos + 1} description",
            )
        with col_up:
            if st.button("↑", key=f"up::{key}", disabled=(pos == 0)):
                order[pos - 1], order[pos] = order[pos], order[pos - 1]
                st.session_state["order"] = order
                st.rerun()
        with col_dn:
            if st.button("↓", key=f"dn::{key}", disabled=(pos == len(order) - 1)):
                order[pos + 1], order[pos] = order[pos], order[pos + 1]
                st.session_state["order"] = order
                st.rerun()

    st.divider()

    items = _load_items(files, order)
    if not items:
        st.error("None of the uploaded files could be read.")
        return

    for layout_key in ("wkw", "gladstone"):
        layout = LAYOUTS[layout_key]
        st.subheader(layout["label"])
        st.caption(
            f"Photos and captions aligned {layout['align_word']}, "
            f"Arial {FONT_SIZE_PT}pt fixed on every run. "
            "Open it, select all, copy, and paste into the report."
        )
        st.download_button(
            "Download .docx",
            data=build_docx(items, layout_key),
            file_name=f"Survey Photo Sheet - {layout['label']}.docx",
            mime=("application/vnd.openxmlformats-officedocument"
                  ".wordprocessingml.document"),
            key=f"dl::{layout_key}",
        )


if __name__ == "__main__":
    main()
